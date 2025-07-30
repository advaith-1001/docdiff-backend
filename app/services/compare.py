from fastapi import UploadFile

from app.models.schemas import ComparisonResult

def compare_texts(text1: str, text2: str) -> ComparisonResult:
    set1 = set(text1.split())
    set2 = set(text2.split())

    common = set1 & set2
    only_in_1 = set1 - set2
    only_in_2 = set2 - set1

    return ComparisonResult(
        common_words=list(common),
        only_in_file1=list(only_in_1),
        only_in_file2=list(only_in_2)
    )


import fitz  # PyMuPDF


async def compare_pdf_formatting(file1: UploadFile, file2: UploadFile):

    file1_bytes = await file1.read()
    file2_bytes = await file2.read()

    doc1 = fitz.open(stream=file1_bytes, filetype="pdf")
    doc2 = fitz.open(stream=file2_bytes, filetype="pdf")
    differences = []

    if len(doc1) != len(doc2):
        differences.append({
            "type": "Page Count Mismatch",
            "location": "Document",
            "value1": f"{len(doc1)} pages",
            "value2": f"{len(doc2)} pages"
        })

    for page_num in range(min(len(doc1), len(doc2))):
        page1 = doc1.load_page(page_num)
        page2 = doc2.load_page(page_num)

        dict1 = page1.get_text("dict")
        dict2 = page2.get_text("dict")

        blocks1 = dict1.get('blocks', [])
        blocks2 = dict2.get('blocks', [])

        if len(blocks1) != len(blocks2):
            differences.append({
                "type": "Structure Mismatch",
                "location": f"Page {page_num + 1}",
                "value1": f"{len(blocks1)} text blocks",
                "value2": f"{len(blocks2)} text blocks"
            })

        for i, (block1, block2) in enumerate(zip(blocks1, blocks2)):
            spans1 = block1.get('lines', [{}])[0].get('spans', [])
            spans2 = block2.get('lines', [{}])[0].get('spans', [])

            if len(spans1) != len(spans2):
                differences.append({
                    "type": "Structure Mismatch",
                    "location": f"Page {page_num + 1}, Block {i + 1}",
                    "value1": f"{len(spans1)} text spans",
                    "value2": f"{len(spans2)} text spans"
                })

    return differences

import docx


def compare_docx_formatting(file1: UploadFile, file2: UploadFile):

    doc1 = docx.Document(file1.file)
    doc2 = docx.Document(file2.file)
    differences = []

    # Check for different number of paragraphs
    if len(doc1.paragraphs) != len(doc2.paragraphs):
        differences.append({
            "type": "Structure Mismatch",
            "location": "Document",
            "value1": f"{len(doc1.paragraphs)} paragraphs",
            "value2": f"{len(doc2.paragraphs)} paragraphs"
        })

    # Iterate through the smaller number of paragraphs to avoid errors
    for i, (p1, p2) in enumerate(zip(doc1.paragraphs, doc2.paragraphs)):
        # Check for different number of runs in a paragraph
        if len(p1.runs) != len(p2.runs):
            differences.append({
                "type": "Structure Mismatch",
                "location": f"Paragraph {i + 1}",
                "value1": f"{len(p1.runs)} text runs",
                "value2": f"{len(p2.runs)} text runs"
            })


        for j, (run1, run2) in enumerate(zip(p1.runs, p2.runs)):
            if run1.bold != run2.bold:
                differences.append({
                    "type": "Bold",
                    "location": f"Paragraph {i + 1}, Text: '{run1.text}'",
                    "value1": run1.bold,
                    "value2": run2.bold
                })

    return differences

